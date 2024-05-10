import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from helpers.utils import reverse_date
from docx.shared import Cm, Pt




def generate_report_file(data: dict) -> dict:
    # Создаем документ
    document = Document()
    ticket_data = data['ticketData']

    # Добавляем абзац с текстом "И.о. директора"
    director_text = document.add_paragraph("И.о. директора")
    director_text.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Выравнивание текста вправо
    director_text.paragraph_format.space_after = Pt(0)  # Устанавливаем нулевой отступ снизу

    # Добавляем абзац с текстом "УрТК НИЯУ МИФИ"
    institution_text = document.add_paragraph("УрТК НИЯУ МИФИ")
    institution_text.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Выравнивание текста вправо
    institution_text.paragraph_format.space_after = Pt(0)  # Устанавливаем нулевой отступ снизу


    # Добавляем абзац с текстом "Тарасову Д.В."
    institution_text = document.add_paragraph("Тарасову Д.В.")
    institution_text.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Выравнивание текста вправо
    institution_text.paragraph_format.space_after = Pt(0)  # Устанавливаем нулевой отступ снизу
    
    # Добавляем абзац с текстом "УрТК НИЯУ МИФИ"
    institution_text = document.add_paragraph("От преподавателя ЦМК №4")
    institution_text.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Выравнивание текста вправо
    institution_text.paragraph_format.space_after = Pt(0)  # Устанавливаем нулевой отступ снизу

    # Добавляем абзац с текстом "УрТК НИЯУ МИФИ"
    institution_text = document.add_paragraph(ticket_data['customer_name'])
    institution_text.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Выравнивание текста вправо
    institution_text.paragraph_format.space_after = Pt(5)  # Устанавливаем нулевой отступ снизу


    # Добавляем абзац с текстом "Заявка системному администратору"
    ticket_title = document.add_paragraph("Заявка системному администратору")
    ticket_title.paragraph_format.space_before = Cm(1.25)

    # Устанавливаем выравнивание текста в центр
    paragraph_format = ticket_title.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Добавляем информацию о тикете
    ticket_title_paragraph = document.add_paragraph()
    ticket_title_paragraph.add_run("Заголовок проблемы: ").bold = True
    ticket_title_paragraph.add_run(ticket_data['problem_title'])

    room_number_paragraph = document.add_paragraph()
    room_number_paragraph.add_run("Номер аудитории: ").bold = True
    room_number_paragraph.add_run(ticket_data['room_number'])

    if (ticket_data['problem_description']):
        description_paragraph = document.add_paragraph()
        description_paragraph.add_run("Описание проблемы: ").bold = True
        description_paragraph.add_run(ticket_data['problem_description'])

    # Добавляем информацию о тикете
    report_period_paragraph = document.add_paragraph()
    report_period_run = report_period_paragraph.add_run("Отчетный период: ")
    report_period_run.bold = True
    report_period_paragraph.add_run(f"с {ticket_data['submission_date']} по {reverse_date(ticket_data['deadline_date'])}")

    performer_name_paragraph = document.add_paragraph()
    performer_name_paragraph.add_run("ФИО исполнителя: ").bold = True
    performer_name_paragraph.add_run(ticket_data['performer_name'])
    
    # Добавляем информацию о задачах
    document.add_heading('Задачи', level=1)
    table = document.add_table(rows=1, cols=2)

    # Заголовки столбцов
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Компьютер'
    hdr_cells[0].paragraphs[0].runs[0].bold = True  # Делаем текст жирным
    hdr_cells[1].text = 'Описание задачи'
    hdr_cells[1].paragraphs[0].runs[0].bold = True  # Делаем текст жирным

    # Устанавливаем ширину первого столбца
    table.columns[0].width = Cm(3)  # Например, установим ширину в 3 сантиметра
    
    # Задачи
    for task in ticket_data['tasks']:
        row_cells = table.add_row().cells
        row_cells[0].text = task['pc_name']
        row_cells[1].text = task['task_description']

    # Добавляем текст с подписью внизу документа
    signature_paragraph = document.add_paragraph("Подпись______________/_______________")
    signature_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Выравниваем текст по правому краю
    signature_paragraph.paragraph_format.space_before = Cm(3)

    # Создаем имя файла на основе ticket_id
    report_filename = f"report_{ticket_data['ticket_id']}.docx"

    # Сохраняем документ
    document.save(os.path.join("reports", report_filename))

    return {"filename": report_filename}

def get_report_file_path(filename: str) -> str | bool:
    report_path = os.path.join("reports", filename)

    if not os.path.exists(report_path):
        return False

    return report_path